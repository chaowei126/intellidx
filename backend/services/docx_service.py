from docx import Document

def generate_docx(content: str, output_path: str):
    doc = Document()
    doc.add_heading('IntelliDX Report', 0)
    doc.add_paragraph(content)
    doc.save(output_path)
    return output_path
